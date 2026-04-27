import xml.etree.ElementTree as ET
import pandas as pd
from docx import Document

import re
from bs4 import BeautifulSoup
import json

xml_file = "wordpress-export.xml"

tree = ET.parse(xml_file)
root = tree.getroot()

ns = {
    'wp': 'http://wordpress.org/export/1.2/',
    'content': 'http://purl.org/rss/1.0/modules/content/',
    'excerpt': 'http://wordpress.org/export/1.2/excerpt/'
}

# -------------------------
# SAFE HELPER FUNCTION
# -------------------------
def get_text(item, tag, ns=None):
    el = item.find(tag, ns) if ns else item.find(tag)
    return el.text.strip() if el is not None and el.text else ""

# Helper to extract meta value for PDF attachment
def get_pdf_attachment(item):
    # Look for <wp:meta_key>_wp_attached_file</wp:meta_key> and get its <wp:meta_value>
    for meta in item.findall('wp:postmeta', ns):
        key = get_text(meta, 'wp:meta_key', ns)
        if key == '_wp_attached_file':
            return get_text(meta, 'wp:meta_value', ns)
    return ""


# -------------------------
# CLEAN CONTENT FUNCTION
# -------------------------
def extract_acf_blocks(raw):
    # Find all Gutenberg/ACF block comments
    # Example: <!-- wp:acf/banner-with-text-and-line-art { ... } /-->
    pattern = r'<!--\s*wp:acf/[^\s]+\s+(\{.*?\})\s*/-->'
    matches = re.findall(pattern, raw, re.DOTALL)
    blocks = []
    for m in matches:
        try:
            block = json.loads(m)
            blocks.append(block)
        except Exception:
            continue
    return blocks

def extract_text_from_acf_block(block):
    # Recursively extract all string values from the block's 'data' dict
    def walk(obj):
        texts = []
        if isinstance(obj, dict):
            for k, v in obj.items():
                if isinstance(v, str):
                    # Only keep non-empty, non-field keys
                    if v.strip() and not k.startswith('_') and not k.endswith('image') and not k.endswith('icon') and not k.endswith('id') and not k.endswith('css') and not k.endswith('class'):
                        texts.append(v.strip())
                elif isinstance(v, (dict, list)):
                    texts.extend(walk(v))
        elif isinstance(obj, list):
            for v in obj:
                texts.extend(walk(v))
        return texts
    if 'data' in block:
        return walk(block['data'])
    return []

def clean_content(content):
    if not content:
        return ""

    content = content.strip()

    # Skip serialized / ACF junk
    if content.startswith("a:") or "s:" in content[:50]:
        return ""

    # Extract all ACF Gutenberg blocks
    blocks = extract_acf_blocks(content)
    all_texts = []
    for block in blocks:
        all_texts.extend(extract_text_from_acf_block(block))

    # Fallback: extract visible HTML text if no blocks found
    if not all_texts:
        soup = BeautifulSoup(content, "html.parser")
        text = soup.get_text(separator="\n")
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    return "\n\n".join(all_texts)


# -------------------------
# MAIN DATA COLLECTION
# -------------------------
data = []



# Extract all <title>, <excerpt:encoded>, <content:encoded>, and PDF attachment fields
for item in root.findall('./channel/item'):
    title = get_text(item, 'title')
    excerpt = get_text(item, 'excerpt:encoded', ns)
    content = get_text(item, 'content:encoded', ns)
    pdf_attachment = get_pdf_attachment(item)

    # Clean up content and excerpt
    clean = clean_content(content)
    excerpt_clean = excerpt.strip() if excerpt else ""

    # Only add row if at least one field is non-empty
    if not (title or clean or excerpt_clean or pdf_attachment):
        continue

    data.append({
        "Title": title,
        "Excerpt": excerpt_clean,
        "Content": clean,
        "PDF Attachment": pdf_attachment
    })


# -------------------------
# EXPORT TO EXCEL
# -------------------------

df = pd.DataFrame(data)
df.to_excel("website_content.xlsx", index=False)



# -------------------------
# EXPORT TO WORD
# -------------------------

doc = Document()
doc.add_heading("Website Content Export", 0)

for row in data:
    doc.add_heading(row["Title"], level=1)
    if row.get("Excerpt"):
        doc.add_paragraph("Excerpt: " + row["Excerpt"])
    if row.get("Content"):
        doc.add_paragraph(row["Content"])
    if row.get("PDF Attachment"):
        doc.add_paragraph("PDF Attachment: " + row["PDF Attachment"])
    doc.add_page_break()

doc.save("website_content.docx")

print("✔ Done - All posts/pages exported successfully")